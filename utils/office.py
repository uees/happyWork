# -- coding: utf-8 -*-
"""
Created on 2015年7月10日

@author: Wan
"""
from docx import Document
from openpyxl import Workbook, load_workbook


class WordTemplate(object):

    def __init__(self, filename):
        self.document = Document(filename)

    def save(self, filename):
        self.document.save(filename)

    def change_paragraph(self, paragraph, filter_format):
        for sformat, content in filter_format.items():
            search_word = "{" + sformat + "}"
            needCheck = True if search_word in paragraph.text else False
            for run in paragraph.runs:
                if search_word in run.text:
                    replaced = run.text.replace(search_word, content)
                    run = run.clear()
                    run.add_text(replaced)
                    needCheck = False
            if needCheck:
                # print sformat, " is not checked in ", paragraph.text
                start_idx = 0
                end_idx = 0
                complete_word = ""
                i = 0
                broken_keys = []
                for run in paragraph.runs:
                    if run.text == "{":
                        start_idx = i
                        complete_word = run.text
                    else:
                        end_idx = i
                        complete_word = complete_word + run.text
                        if run.text == "}":
                            replaced = complete_word.replace(search_word, content)
                            broken_keys.append((start_idx, end_idx, replaced))
                    # print ":::", run.text
                    i = i + 1

                for (start_idx, end_idx, replaced) in broken_keys:
                    for run_idx, run in enumerate(paragraph.runs):
                        if start_idx <= run_idx and run_idx <= end_idx:
                            run = run.clear()
                        if run_idx == end_idx:
                            # print "))", replaced
                            run.add_text(replaced)

    def replace(self, filter_format):
        """ code is very rubbish, but it Cross-platform """
        for paragraph in self.document.paragraphs:
            self.change_paragraph(paragraph, filter_format)

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.change_paragraph(paragraph, filter_format)


class XlsOperator(object):
    """
    Some convenience methods for Excel documents accessed
    through OpenPyXl.
    """

    def __init__(self, filename=None):
        '''
        if filename is None, create a new file
        else, open an exsiting one
        '''
        if filename:
            self.filename = filename
            self.wb = load_workbook(self.filename)
        else:
            self.filename = 'newfile.xlsx'
            self.wb = Workbook()

    def save(self, filename=None):
        """
        if filename is None, save the openning file
        else save as another file used the given name.
        This operation will overwrite existing files without warning.
        """
        if filename:
            self.filename = filename
        self.wb.save(filename=self.filename)

    def select(self, sheet=None):
        """ select and return a sheet """
        if sheet:
            self.ws = self.wb.get_sheet_by_name(sheet)
        else:
            self.ws = self.wb.active
        return self.ws

    def create_sheet(self, name=None):
        """ create and return a new sheet """
        ws = self.wb.create_sheet(title=name)
        return ws

    def get_cell(self, point):
        """
        Get value of one cell
        @point: string cell's coordinate. egg, "A4"
        """
        return self.ws.cell(point).value

    def set_cell(self, point, value):
        """
        Set value of one cell
        """
        self.ws.cell(point).value = value

    def get_range(self, point1, point2):
        """
        ws.iter_rows('A1:C2')
        """
        return self.ws.iter_rows('%s:%s' % (point1, point2))

    def count_rows(self, sheet):
        """
        return the number of the sheet rows.
        """
        ws = self.wb.get_sheet_by_name(sheet)
        return ws.max_row  # len(ws.rows)
